from docx import Document
# from docx.shared import Imches I cannot import the function Inches in my cas
# this function help to reshape the foto
document = Document()
#name = "Tondjua"
#phone_number = "017628059114"
#email = "boris.tondjua@gmail.com"
#document.add_paragraph( name + "|" + phone_number + "|" + email)

# my picture
document.add_picture("boris.png.PNG")
# my name phone number and adress
name = input("What is your name ?")
phone_number = input("What is your phone_number ?")
email = input("What is your email ?")
document.add_paragraph( name + "|" + phone_number + "|" + email)

# about me
document.add_heading("About me")
about_me = input("tell about yourself")
document.add_paragraph(about_me)
# Work experiences
document.add_paragraph("work experiences")
p = document.add_paragraph()
compagny = input("Enter compagny")
from_date = input("from date")
to_date = input("To date")
p.add_run(compagny + " ").bold = True
p.add_run(from_date + " " + to_date + "\n")
experience_details = input(
    "Descriebe your experience at " + compagny)
# more experiences
while True:
            has_more_experiences = input(
                "Do you have more experiences ? Yes or No ")
            if has_more_experiences.lower() == "yes":
                p = document.add_paragraph()
                compagny = input("Enter compagny")
                from_date = input("from date")
                to_date = input("To date")
                p.add_run(compagny + " ").bold = True
                p.add_run(from_date + " " + to_date + "\n")
                experience_details = input(
                    "Descriebe your experience at " + compagny + " gg")
                # more experiences
                p.add_run(experience_details)
            else:
                break
# Sills
document.add_heading("skills")
skill = input("Enter skills")
p = document.add_paragraph(skill)
#p.style = "list Bullet" why I receve KeyError: "no style with name 'list Bullet'"

while True:
            has_more_skills = input("Do you have more skills ? yes or No")
            if has_more_skills.lower() == "yes":
                skill = input("Enter skill")
                p = document.add_paragraph(skill)
                #p.style = "list Bullet"
            else:
                break

# footer to our document
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using BorisCode"

# import  pyttsx3 With this Packages I cann year the speech from my computer but i cannot inst it

document.save("cv_boris.dock")